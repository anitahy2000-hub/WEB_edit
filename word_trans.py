#!/usr/bin/env python3
"""
Word (.docx) 文档转 HTML 转换器
- 支持文本、标题、粗体、斜体、下划线、颜色等格式
- 图片自动转为 Base64 内嵌
- 支持表格、列表（有序/无序）
- 生成可直接打开的独立 HTML 文件

依赖安装：
    pip install python-docx
"""

import base64
import os
import sys
import zipfile
import re
from io import BytesIO
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
except ImportError:
    print("❌ 缺少依赖，请先运行：pip install python-docx")
    sys.exit(1)

# ── 颜色与对齐映射 ──────────────────────────────────────────────
ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT:    "left",
    WD_ALIGN_PARAGRAPH.CENTER:  "center",
    WD_ALIGN_PARAGRAPH.RIGHT:   "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

# ── 辅助：提取 docx 内图片（rId → base64 data URI）─────────────
def extract_images(docx_path: str) -> dict:
    """从 docx ZIP 包中提取所有图片，返回 {rId: 'data:image/xxx;base64,...'}"""
    images = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        # 解析 word/_rels/document.xml.rels
        rels_path = "word/_rels/document.xml.rels"
        if rels_path not in z.namelist():
            return images

        rels_xml = z.read(rels_path).decode("utf-8")
        # 找出所有图片关系
        pattern = r'Id="([^"]+)"[^>]+Type="[^"]*image[^"]*"[^>]+Target="([^"]+)"'
        for rid, target in re.findall(pattern, rels_xml):
            # target 可能是 media/imageX.xxx 或 ../media/...
            target = target.lstrip("../")
            full_path = f"word/{target}" if not target.startswith("word/") else target
            if full_path in z.namelist():
                data = z.read(full_path)
                ext = Path(target).suffix.lower().lstrip(".")
                mime = {
                    "jpg": "jpeg", "jpeg": "jpeg", "png": "png",
                    "gif": "gif",  "bmp": "bmp",  "webp": "webp",
                    "svg": "svg+xml", "tiff": "tiff", "emf": "x-emf",
                    "wmf": "x-wmf",
                }.get(ext, ext)
                b64 = base64.b64encode(data).decode("ascii")
                images[rid] = f"data:image/{mime};base64,{b64}"
    return images


# ── 辅助：run 的内联样式 ────────────────────────────────────────
def run_style(run) -> str:
    styles = []
    if run.bold:
        styles.append("font-weight:bold")
    if run.italic:
        styles.append("font-style:italic")
    if run.underline:
        styles.append("text-decoration:underline")
    if run.font.strike:
        styles.append("text-decoration:line-through")
    if run.font.color and run.font.color.type is not None:
        try:
            rgb: RGBColor = run.font.color.rgb
            styles.append(f"color:#{rgb}")
        except Exception:
            pass
    if run.font.size:
        pt = run.font.size.pt
        styles.append(f"font-size:{pt:.1f}pt")
    return ";".join(styles)


# ── 辅助：解析段落内的 run（含图片）──────────────────────────────
def parse_runs(para, images: dict) -> str:
    """将段落中的 run 转为 HTML 片段，自动处理图片"""
    html = ""
    for child in para._p:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        # 普通文字 run
        if tag == "r":
            # 先检查是否包含图片
            drawing = child.find(qn("w:drawing"))
            pict    = child.find(qn("w:pict"))

            if drawing is not None or pict is not None:
                # 提取 rId —— r:embed 对应的完整命名空间 key
                REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                rid = None
                for elem in child.iter():
                    etag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                    if etag == "blip":
                        # 正确的关系命名空间
                        rid = elem.get(f"{{{REL_NS}}}embed")
                        if not rid:
                            # 兜底：遍历全部属性，找任意以 }embed 结尾的键
                            for attr_key, attr_val in elem.attrib.items():
                                if attr_key.endswith("}embed") or attr_key == "embed":
                                    rid = attr_val
                                    break
                        break
                    if etag == "imagedata":
                        rid = (elem.get(f"{{{REL_NS}}}id")
                               or elem.get(f"{{{REL_NS}}}href"))
                        break

                if rid and rid in images:
                    # 读取图片名称
                    img_name = "image"
                    for elem in child.iter():
                        etag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                        if etag in ("docPr", "cNvPr"):
                            n = elem.get("name") or elem.get("descr") or ""
                            if n:
                                img_name = n
                                break
                    # 按旧式 HTML 格式输出：mime 单独一行，base64 每 76 字符换行
                    data_uri = images[rid]
                    header, b64data = data_uri.split(",", 1)
                    mime = header.split(";")[0]   # e.g. "data:image/png"
                    chunks = [b64data[i:i+76] for i in range(0, len(b64data), 76)]
                    b64_lines = "\n       ".join(chunks)
                    src_value = "\n      " + mime + ";\n      " + mime + ";base64," + b64_lines + "\n      "
                    html += '<img src="' + src_value + '" alt="' + img_name + '" />'
                else:
                    html += '[图片未找到]'
            else:
                # 文字 run
                text = ""
                for t in child.findall(qn("w:t")):
                    text += t.text or ""
                for br in child.findall(qn("w:br")):
                    text += "\n"

                if not text:
                    continue

                text_html = (text.replace("&", "&amp;")
                                 .replace("<", "&lt;")
                                 .replace(">", "&gt;")
                                 .replace("\n", "<br/>"))

                # 构造临时 run 对象仅用于样式判断
                from docx.text.run import Run
                try:
                    r_obj = Run(child, para)
                    style = run_style(r_obj)
                except Exception:
                    style = ""

                if style:
                    html += f'<span style="{style}">{text_html}</span>'
                else:
                    html += text_html

        # 超链接
        elif tag == "hyperlink":
            link_url = ""
            for attr, val in child.attrib.items():
                if "id" in attr.lower():
                    link_url = val
                    break
            inner = ""
            for r in child.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    inner += (t.text or "").replace("&","&amp;").replace("<","&lt;")
            if link_url:
                html += f'<a href="{link_url}">{inner}</a>'
            else:
                html += inner

    return html


# ── 辅助：段落对应的 HTML 标签 ──────────────────────────────────
def heading_tag(style_name: str):
    """根据样式名返回标题标签，None 表示普通段落"""
    s = style_name.lower()
    for i in range(1, 10):
        if f"heading {i}" in s or f"标题 {i}" in s or f"标题{i}" in s:
            return f"h{i}"
    return None


# ── 核心：段落 → HTML ───────────────────────────────────────────
def para_to_html(para, images: dict, list_state: dict) -> str:
    style_name = para.style.name if para.style else ""
    h_tag = heading_tag(style_name)

    inner = parse_runs(para, images)

    # 对齐
    align = ALIGN_MAP.get(para.alignment, "")
    align_style = f"text-align:{align};" if align else ""

    # 段落间距
    space_before = ""
    space_after  = ""
    try:
        if para.paragraph_format.space_before:
            space_before = f"margin-top:{para.paragraph_format.space_before.pt:.0f}pt;"
        if para.paragraph_format.space_after:
            space_after = f"margin-bottom:{para.paragraph_format.space_after.pt:.0f}pt;"
    except Exception:
        pass

    p_style = align_style + space_before + space_after

    # 列表
    is_list = False
    list_level = 0
    list_type  = "ul"
    num_pr = para._p.find(qn("w:pPr"))
    if num_pr is not None:
        num_id_elem = num_pr.find(qn("w:numId"))
        ilvl_elem   = num_pr.find(qn("w:ilvl"))
        if num_id_elem is not None:
            is_list = True
            list_level = int(ilvl_elem.get(qn("w:val"), "0")) if ilvl_elem is not None else 0

    # 标题
    if h_tag:
        list_state["open"] = False
        tag_style = p_style
        return f'<{h_tag} style="{tag_style}">{inner}</{h_tag}>\n'

    # 列表项
    if is_list:
        indent = list_level * 20
        li_style = f"margin-left:{indent}px;" + p_style
        return f'<li style="{li_style}">{inner}</li>\n'

    # 空段落
    if not inner.strip():
        return '<p style="margin:4px 0;">&nbsp;</p>\n'

    return f'<p style="{p_style}margin:6px 0;">{inner}</p>\n'


# ── 核心：表格 → HTML ───────────────────────────────────────────
def table_to_html(table, images: dict) -> str:
    html = '<table style="border-collapse:collapse;width:100%;margin:12px 0;">\n'
    for row in table.rows:
        html += "  <tr>\n"
        for cell in row.cells:
            # 表头判断：首行
            td = "th" if row == table.rows[0] else "td"
            bg = ' style="background:#f2f2f2;font-weight:bold;"' if td == "th" else ' style=""'
            td_style = 'border:1px solid #ccc;padding:6px 10px;vertical-align:top;'
            cell_html = ""
            for para in cell.paragraphs:
                cell_html += parse_runs(para, images)
                cell_html += "<br/>"
            cell_html = cell_html.rstrip("<br/>")
            html += f'    <{td} style="{td_style}">{cell_html}</{td}>\n'
        html += "  </tr>\n"
    html += "</table>\n"
    return html


# ── 主转换函数 ──────────────────────────────────────────────────
def docx_to_html(docx_path: str, output_path: str | None = None) -> str:
    """
    将 docx 文件转换为 HTML 字符串，并可选写出到文件。
    """
    docx_path = str(docx_path)
    doc    = Document(docx_path)
    images = extract_images(docx_path)

    title = Path(docx_path).stem

    body_parts = []
    list_state = {"open": False, "type": "ul"}

    # 遍历文档 body 中的顶层元素（段落 + 表格）
    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    para_idx  = 0
    table_idx = 0

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = doc.paragraphs[para_idx] if para_idx < len(doc.paragraphs) else None
            para_idx += 1
            if para is None:
                continue
            html_chunk = para_to_html(para, images, list_state)
            body_parts.append(html_chunk)

        elif tag == "tbl":
            if table_idx < len(doc.tables):
                body_parts.append(table_to_html(doc.tables[table_idx], images))
                table_idx += 1

    body_html = "".join(body_parts)

    # 完整 HTML 模板
    html = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8"/>
  <title>{title}</title>
</head>
<body>
  <div>
{body_html}
  </div>
</body>
</html>
"""


    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)
        print(f"✅ 转换完成：{output_path}")
    return html


# ── 命令行入口 ──────────────────────────────────────────────────
def main():
    if len(sys.argv) < 2:
        print("用法：python docx_to_html.py <input.docx> [output.html]")
        print("示例：python docx_to_html.py report.docx report.html")
        sys.exit(1)

    input_file  = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) >= 3 else Path(input_file).stem + ".html"

    if not os.path.exists(input_file):
        print(f"❌ 文件不存在：{input_file}")
        sys.exit(1)

    docx_to_html(input_file, output_file)


if __name__ == "__main__":
    main()